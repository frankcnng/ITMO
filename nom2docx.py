import json
import argparse
import sys

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn

from datetime import datetime

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

# Indent in twip = 566/cm
def indent_table(table, indent):
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(indent))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)

def formatTable(t):
    table.autofit=False
    table.allow_autofit=False
    r=0
    for row in table.rows:
        r=r+1
        for cell in row.cells:
            if r==1:
                set_cell_border(cell,
                    top={"sz": 4, "val": "single", "color": "#000000", "space": "0"},
                    bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"})
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                if r==1:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing_rule=WD_LINE_SPACING.EXACTLY
                    paragraph.paragraph_format.line_spacing=Pt(9)
                else:
                    paragraph.paragraph_format.space_after = Pt(4)
                    paragraph.paragraph_format.line_spacing_rule=WD_LINE_SPACING.AT_LEAST
                    paragraph.paragraph_format.line_spacing=Pt(12)
                for run in paragraph.runs:
                    font = run.font
                    if r==1:
                        font.size=Pt(8.0)
                        font.italic=True
                    else:
                        font.size=Pt(10.0)
    # Bottom line
    for cell in row.cells:
        set_cell_border(cell,
            bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"})

    col = table.columns[0] 
    col.width=Cm(3.0)
    col = table.columns[1] 
    col.width=Cm(10.0)
    for cell in table.columns[0].cells:
        cell.width = Cm(3.0)
    for cell in table.columns[1].cells:
        cell.width = Cm(10.0)
    
    set_repeat_table_header(table.rows[0])
    indent_table(table, Cm(2).twips)

    #Add a para between tables
    p=doc.add_paragraph("")
    p.paragraph_format.space_before=Pt(0)
    p.paragraph_format.space_after=Pt(6)

#Main
parser=argparse.ArgumentParser(prog="nom2docx",description="Convert common nomenclature JSON to DOCX")
parser.add_argument("-n", "--nomenclature", dest="nomenclature",
                    help="The common nomenclature file to be converted (JSON)",
                    default="CommonNomenclature.json")
parser.add_argument("-t", "--template", dest="template",
                    help="The common nomenclature template file (DOCX)",
                    default="CommonNomenclatureTemplate.docx")

parser.add_argument("-o", "--output", dest="output",
                    help="The common nomenclature output file (DOCX)",
                    default="CommonNomenclature.docx")
args = parser.parse_args()

#args.nomenclature=r"C:/Users/Halleux/OneDrive - United Nations Framework Convention on Climate Change/Official/CommonNomenclature.json"
#args.template=r"C:/Users/Halleux/OneDrive - United Nations Framework Convention on Climate Change/Official/CommonNomenclatureTemplate.docx"
#args.output=r"C:/Users/Halleux/OneDrive - United Nations Framework Convention on Climate Change/Official/CommonNomenclature.docx"

with open(args.nomenclature,"r") as jsonfile:
    contents=json.loads(jsonfile.read())

doc=Document(args.template)

tnb=0
for n in contents['CommonNomenclature']['contents']:
    # Check availability of definition
    try:
        _ = contents[n]
    except KeyError as ke:
        print(ke,"is defined in contents but is missing a definition")
        print()
        print("Halting processing")
        sys.exit(1)

    tnb=tnb+1
    # Metadata
    # Table caption and title
    p=doc.add_paragraph()
    if 'definedValues' in contents[n]:
        run=p.add_run("Table "+str(tnb)+"(a)\n")
    else:
        run=p.add_run("Table "+str(tnb)+"\n")
    run=p.add_run(contents[n]['attribute']+": metadata")
    run.bold=True
    p.paragraph_format.space_before=Pt(0)
    p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.left_indent=Cm(2)
    p.paragraph_format.keep_with_next=True
    # Table itself
    table=doc.add_table(rows=8,cols=2)
    row=table.rows[0].cells 
    row[0].text = 'Element'
    row[1].text = 'Value'
    row=table.rows[1].cells 
    row[0].text='Description'
    row[1].text=contents[n]['description']
    row=table.rows[2].cells 
    row[0].text='Mandate(s)'
    mt=''
    for m in contents[n]['requiredBy']: mt=mt+m+"\n"
    row[1].text=mt[:-1]
    row=table.rows[3].cells 
    row[0].text='Type'
    row[1].text=contents[n]['type']
    row=table.rows[4].cells
    row[0].text='Required'
    row[1].text=contents[n]['required']
    row=table.rows[5].cells
    row[0].text='Naming in AEF'
    aef=''
    for naef in contents[n]['AEF fields']: aef=aef+naef+"\n"
    row[1].text=aef[:-1]
    row=table.rows[6].cells
    row[0].text='Technical name'
    row[1].text=n
    
    formatTable(table)

    if 'definedValues' not in contents[n]:
        row=table.rows[7].cells
        row[0].text='List of values'
        row[1].text='Not available'
    else:
        #List of values
        row=table.rows[7].cells
        row[0].text='List of values'
        row[1].text='See table '+str(tnb)+"(b) below"
        # Table caption and title
        p=doc.add_paragraph()
        run=p.add_run("Table "+str(tnb)+"(b)\n")
        run=p.add_run(contents[n]['attribute']+": list of values and descriptions")
        run.bold=True
        p.paragraph_format.space_before=Pt(0)
        p.paragraph_format.space_after=Pt(6)
        p.paragraph_format.left_indent=Cm(2)
        # Table itself
        table=doc.add_table(rows=1,cols=2)
        row=table.rows[0].cells 
        row[0].text = 'Value'
        row[1].text = 'Description'
        for v in contents[n]['definedValues']:
            row = table.add_row().cells
            row[0].text=v['value']
            row[1].text=v['description']
        formatTable(table)

doc.save(args.output)
